import PIL.Image #for image processing
import pytesseract # for photo to text
import os # for fliepaths and os things
import docx
from docx import Document #to write to .docs files   INSTALL python-docx
import unicodedata #remove unicode to transfer to .docx
import re # REGEX LET'S GO

#https://github.com/UB-Mannheim/tesseract/wiki


#add a user input to continue
def inputToContinue(question):
    ask = input(question)
    return ask

#create a relative file path to reference
userprofile = os.environ['USERPROFILE']
AppDataPath = os.path.join(userprofile, 'AppData')
print("userprofile is: " + userprofile)

#set the file location of tesseract
pytesseract.pytesseract.tesseract_cmd = AppDataPath + "\\local\\programs\\Tesseract-OCR\\tesseract.exe"

#write to .docx prerequesites
document = Document()

##############################################################################################
##############################################################################################


def processImage():
    #ask folder path
    imgPath = askopenfilename(title='select your translations document') # shows dialog box and return the path
    print(imgPath)
    dirLabel.configure(text="the translation image folderpath is: " + imgPath)

    #process the image
    text = pytesseract.image_to_string(PIL.Image.open(imgPath), lang = 'eng')
    print(text)
    
    with open('transl.docx', mode ='w'):
        #re.match("[^a-zA-Z0-9 ,.'\"\n:();]")
        formattedText = re.sub("[^ -~\"]", " ", text)
        print(formattedText)
        document.add_paragraph(formattedText)
        document.save('transl.docx')


def doRevisions():
    #code asks you to edit before you continue
    #inputToContinue(col.white + "please press enter to make revisions to the text file")

    #run the text file to edit
    os.startfile(os.path.dirname(__file__) + "\\transl.docx")
    os.startfile(os.path.dirname(__file__) + "\\transl.png")

    #asks you to enter once you are done editing
    #inputToContinue(col.white + "please press enter once you are done with revisions to the document")


def translate():
    # Opening MS word document 
    doc = docx.Document('transl.docx')

    
    #Filtering iltalicized text 
    Italic_Text = '' 
    for p in doc.paragraphs: 
        for run in p.runs: 
            if run.italic: 
                Italic_Text = Italic_Text + run.text + " "
    
    #Splitting text to a list of words 
    Italic_Words =list(filter(None,Italic_Text.split(' '))) 
    print(Italic_Words)

    #Counting the number of italicized words 
    counter = len(Italic_Words)
    print(counter)
    """
    #Filtering iltalicized text 
    highlightedText = '' 
    for p in doc.paragraphs: 
        for run in p.runs: 
            if run.italic: 
                highlightedText = highlightedText + run.text + " "
                """

##################################################################################################
##################################################################################################
import tkinter
from tkinter import Tk, Label, Button
from tkinter.filedialog import askopenfilename

#set so you can call tk.whatever
tk = Tk()

#headers for the thing
tk.title("Welcome to translation station")
tk.geometry('600x600')

#what to do when directory button is clicked
def dirButton():
    global folderpath
    folderpath = askopenfilename(title='select your translations image to convert it to a word doc')
    print(folderpath)
    dirLabel.configure(text="the translation image folderpath is: " + folderpath)
    dirBtn.grid(column=0, row=0)
    return folderpath


#the label where the directory will go
dirLabel = Label(tk)   #, text="Hello"
dirLabel.grid(column=0, row=1)

#button for photo to text to word doc button
dirBtn = Button(tk, text="Select your translations image file path to convert it to text", command=processImage)
dirBtn.grid(column=0, row=0)

#button for opening the word docs
docBtn = Button(tk, text="Open the word document to highlight your stem words and do revisions", command=doRevisions)
docBtn.grid(column=0, row=4)

#button to translate
translBtn = Button(tk, text="convert the highlighted words to stems", command=translate)
translBtn.grid(column=0, row=6)


#run the sucker
tk.mainloop()










"""
def writeToDoc(writeBackup):
    #write the text to 'transl.docx' in this directory
    try:
        document.add_paragraph('Translations:')
    except IOError:
            print(col.red + IOError)
    else:
        with open('transl.docx', mode ='w'):
            #re.match("[a-zA-Z0-9 ,.'\"\n:();]")
            formattedText = re.sub("[^a-zA-Z0-9 ,.'\"\n]", " ", text)
            print(col.white + formattedText)
            document.add_paragraph(formattedText)
            document.save('transl.docx')

        if writeBackup == True:
            with open('translBackup.docx', mode ='w'):
                formattedText = re.sub("[^a-zA-Z0-9 ,.'\"\n]", " ", text)
                print(col.white + formattedText)
                document.add_paragraph(formattedText)
                document.save('translBackup.docx')

            #write to the console
            print(col.green + "text file written to " + os.path.dirname(__file__) + "\\transl.docx" )
"""